#!/usr/bin/env python3
"""
build_protocol.py — 잠긴 prereg.json + narrative.json(prose 오버레이) + search_log.json(인용)을
기관 표준 IRB 연구계획서(.docx)로 결정론적 렌더링.

이 스크립트는 anthropic-skills:research-protocol-writer(앱 내장 시스템 스킬) 의존을 제거하기 위한
자체 생성기다. python-docx만 있으면 외부 스킬 없이 단독 동작한다.

SSOT 원칙: 가설·분석·결과변수 등 구조적 필드는 prereg.json이 유일 출처.
narrative.json은 prose(배경, 설계 서술 등)만 덧입힌다. prereg와 충돌하는 narrative 구조 필드는 무시.

출력(모두 --outdir 기준):
  research_protocol.docx          IRB 제출용 한국어 계획서
  protocol_content.resolved.json  렌더링에 쓰인 병합 결과(자기검증·인용검증 입력)
  irb_metadata.json               IRB 추적 메타데이터(linked_prereg_hash 포함)

사용:
  python build_protocol.py --prereg prereg.json --narrative narrative.json \
    --refs search_log.json --profile researcher_profile.json --outdir phase3_protocol
"""
import argparse, json, os, sys, datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    sys.exit("python-docx 미설치: pip install python-docx")

# ---------- 로딩 ----------
def load(path, default=None):
    if not path or not os.path.exists(path):
        return default
    with open(path, encoding="utf-8") as f:
        return json.load(f)

def get(d, path, default=None):
    cur=d
    for k in path.split("."):
        if isinstance(cur,dict) and k in cur: cur=cur[k]
        else: return default
    return cur

# ---------- 병합: prereg(구조) + narrative(prose) + profile(메타) ----------
def resolve(prereg, narrative, profile, refs):
    n=narrative or {}; p=profile or {}
    h=prereg.get("hypothesis",{}); ap=prereg.get("analysis_plan",{}); dp=prereg.get("data_provenance",{})
    pi = n.get("pi") or p.get("pi") or {}
    op=h.get("outcome_primary",{}) or {}
    osec=h.get("outcomes_secondary",[]) or []
    # 결과변수 텍스트화
    def outcome_str(o):
        if isinstance(o,dict):
            parts=[o.get("name","")]
            if o.get("definition"): parts.append(f"({o['definition']})")
            if o.get("timepoint"): parts.append(f"[{o['timepoint']}]")
            return " ".join(x for x in parts if x).strip()
        return str(o)
    content={
        "title_ko": n.get("title_ko") or prereg.get("project","연구계획서"),
        "title_en": n.get("title_en") or "",
        "institution": n.get("institution") or p.get("institution") or "(소속 기관)",
        "pi": {
            "name": pi.get("name") or "(연구책임자)",
            "title": pi.get("title") or "",
            "dept": pi.get("dept") or "",
            "email": pi.get("email") or "",
        },
        "co_investigators": n.get("co_investigators") or p.get("co_investigators") or [],
        "version": str(n.get("version") or prereg.get("version") or "1.0"),
        "date": n.get("date") or "",
        "study_period": n.get("study_period") or "",
        "design": h.get("design",""),
        "design_narrative": n.get("design_narrative") or "",
        "background": n.get("background") or "",
        "objectives": n.get("objectives") or {
            "primary": op.get("name",""), "secondary":[outcome_str(o) for o in osec], "expected":""},
        "subjects": n.get("subjects") or {"inclusion":[h.get("population","")], "exclusion":[]},
        "population": h.get("population",""),
        "exposure": h.get("exposure",""),
        "comparator": h.get("comparator",""),
        "outcome_primary": outcome_str(op),
        "outcomes_secondary": [outcome_str(o) for o in osec],
        "data_collection": n.get("data_collection") or "",
        "data_provenance": dp,
        "ai_algorithm": n.get("ai_algorithm"),
        "statistics": {
            "primary": ap.get("primary_method",""),
            "covariates": ap.get("covariates",[]),
            "sensitivity": ap.get("sensitivity",[]),
            "missing": ap.get("missing_handling",""),
            "multiplicity": ap.get("multiple_comparisons",{}),
            "software": ap.get("software",[]),
            "extra": n.get("statistics_extra") or "",
        },
        "sample_size": n.get("sample_size") or "",
        "ethics": n.get("ethics") or {},
        "references": n.get("references") or refs_to_list(refs),
        "study_type": n.get("study_type") or infer_study_type(h.get("design","")),
        "outline": n.get("outline") or {},
    }
    return content

def infer_study_type(design):
    d=(design or "").lower()
    if any(k in d for k in ["rct","randomized","trial","무작위","임상시험"]): return "trial"
    return "observational"

def refs_to_list(refs):
    """search_log.json을 [{n,text,pmid,doi}] 로 정규화 (없으면 빈 리스트)."""
    if not refs: return []
    items = refs.get("results") if isinstance(refs,dict) else refs
    if not isinstance(items,list): return []
    out=[]
    for i,it in enumerate(items,1):
        if not isinstance(it,dict): continue
        txt=it.get("citation") or it.get("vancouver") or it.get("text") or _fmt_ref(it)
        out.append({"n":i,"text":txt,"pmid":str(it.get("pmid") or ""),"doi":str(it.get("doi") or "")})
    return out

def _fmt_ref(it):
    a=it.get("authors") or ""
    if isinstance(a,list): a=", ".join(a[:3])+(" et al" if len(a)>3 else "")
    return f"{a}. {it.get('title','')}. {it.get('journal','')}. {it.get('year','')}.".strip()

# ---------- docx 헬퍼 ----------
def setup_styles(doc):
    st=doc.styles["Normal"]; st.font.name="맑은 고딕"; st.font.size=Pt(10.5)
    try:
        st._element.rPr.rFonts.set(__import__("docx.oxml.ns",fromlist=["qn"]).qn("w:eastAsia"),"맑은 고딕")
    except Exception: pass

def h1(doc,t):
    p=doc.add_heading(t,level=1);
    for r in p.runs: r.font.color.rgb=RGBColor(0,0,0)
    return p
def h2(doc,t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=RGBColor(0,0,0)
    return p
def para(doc,t="",bold=False,align=None,size=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold
    if size: r.font.size=Pt(size)
    if align=="center": p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    return p
def bullets(doc,items):
    for it in items:
        if it is None or str(it).strip()=="": continue
        doc.add_paragraph(str(it),style="List Bullet")

# ---------- 렌더링 ----------
def pi_line(c):
    """책임연구자 한 줄. 비어있는 소속/과/직위는 깔끔히 생략."""
    pi=c["pi"]
    parts=[c.get("institution","").strip(), pi.get("dept","").strip(), pi.get("name","").strip()]
    s=" ".join(p for p in parts if p)
    if pi.get("title","").strip(): s+=f" ({pi['title'].strip()})"
    return s

def render(content, outpath):
    doc=Document(); setup_styles(doc)
    c=content
    # === 표지 ===
    para(doc,"연구계획서",bold=True,align="center",size=20)
    para(doc)
    para(doc,c["title_ko"],bold=True,align="center",size=14)
    if c["title_en"]: para(doc,c["title_en"],align="center",size=11)
    para(doc)
    pi=c["pi"]
    para(doc,f"책임 연구자: {pi_line(c)}",align="center")
    if c["date"]: para(doc,c["date"],align="center")
    para(doc,f"Ver. {c['version']}",align="center")
    doc.add_page_break()

    # === Protocol Outline (요약 표) ===
    h1(doc,"Protocol Outline")
    outline=c.get("outline",{})
    rows=[
        ("연 구 제 목", c["title_ko"]+(("\n"+c["title_en"]) if c["title_en"] else "")),
        ("연 구 목 적", outline.get("purpose") or get(c,"objectives.primary","")),
        ("연 구 기 관", c["institution"]),
        ("연구책임자", pi_line(c)+(f" | {pi['email']}" if pi['email'] else "")),
        ("공동연구자", ", ".join(c["co_investigators"]) if c["co_investigators"] else "해당 없음"),
        ("연 구 대 상", outline.get("subjects") or c["population"]),
        ("연 구 기 간", c["study_period"] or "IRB 승인/면제 확인 후 ~ (연구자 기입)"),
        ("연 구 방 법", outline.get("methods") or c["statistics"]["primary"]),
        ("기대효과 및 예상결과", outline.get("expected") or get(c,"objectives.expected","")),
    ]
    tbl=doc.add_table(rows=0,cols=2); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for label,val in rows:
        cells=tbl.add_row().cells
        cells[0].width=Cm(4); cells[1].width=Cm(13)
        cells[0].paragraphs[0].add_run(label).bold=True
        cells[1].text=str(val or "")
    doc.add_page_break()

    # === 1. 연구제목 ===
    h1(doc,"1. 연구제목")
    para(doc,f"(국문) {c['title_ko']}")
    if c["title_en"]: para(doc,f"(영문) {c['title_en']}")

    # === 2. 연구 배경 및 목적 ===
    h1(doc,"2. 연구 배경 및 목적")
    if c["background"]:
        for blk in str(c["background"]).split("\n"):
            if blk.strip(): para(doc,blk.strip())
    else:
        para(doc,"[연구 배경 narrative 미작성 — narrative.json의 background 필드를 채우십시오. "
                 "Citation Grounding: 모든 인용은 search_log.json의 PMID 동반 문헌만 사용.]")

    # === 3. 연구목표 및 기대효과 ===
    h1(doc,"3. 연구목표 및 기대효과")
    obj=c["objectives"]
    para(doc,"1차 목표 (Primary objective):",bold=True)
    para(doc,obj.get("primary",""))
    if obj.get("secondary"):
        para(doc,"2차 목표 (Secondary objectives):",bold=True)
        bullets(doc,obj["secondary"])
    if obj.get("expected"):
        para(doc,"기대효과:",bold=True); para(doc,obj["expected"])

    # === 4. 예상 연구기간 ===
    h1(doc,"4. 예상 연구기간")
    para(doc,c["study_period"] or "[연구 기간 미기입 — narrative.json의 study_period]")

    # === 5. 연구 내용 및 방법 ===
    h1(doc,"5. 연구 내용 및 방법")
    h2(doc,"1) 연구 설계")
    para(doc,c["design_narrative"] or f"본 연구는 {c['design']} 설계로 수행한다.")
    h2(doc,"2) 연구 방법")
    para(doc,"연구대상",bold=True)
    sub=c["subjects"]
    if sub.get("inclusion"):
        para(doc,"선정 기준:",bold=True); bullets(doc,sub["inclusion"])
    if sub.get("exclusion"):
        para(doc,"제외 기준:",bold=True); bullets(doc,sub["exclusion"])
    para(doc,"노출/중재 및 비교군",bold=True)
    para(doc,f"노출(중재): {c['exposure']}")
    para(doc,f"비교군: {c['comparator']}")
    para(doc,"결과 변수",bold=True)
    para(doc,f"1차 결과변수: {c['outcome_primary']}")
    if c["outcomes_secondary"]:
        para(doc,"2차 결과변수:"); bullets(doc,c["outcomes_secondary"])
    para(doc,"데이터 수집 및 전처리",bold=True)
    dp=c["data_provenance"]
    dc=c["data_collection"] or (f"자료원: {dp.get('source','(기입)')}. "
        f"분석 대상 N={dp.get('n_records','(기입)')}. 자료사전: {dp.get('data_dictionary_path','(기입)')}.")
    para(doc,dc)
    if c.get("ai_algorithm"):
        para(doc,"인공지능 알고리즘",bold=True); para(doc,c["ai_algorithm"])
    # 통계 분석
    para(doc,"통계 분석",bold=True)
    stt=c["statistics"]
    para(doc,f"1차 분석: {stt['primary']}")
    if stt.get("covariates"):
        para(doc,f"공변량: {', '.join(stt['covariates'])}")
    if stt.get("sensitivity"):
        para(doc,"민감도/하위군 분석(사전 명시):"); bullets(doc,stt["sensitivity"])
    if stt.get("missing"):
        para(doc,f"결측치 처리: {stt['missing']}")
    if stt.get("multiplicity"):
        mc=stt["multiplicity"]
        if isinstance(mc,dict) and mc:
            para(doc,"다중 검정 보정: "+"; ".join(f"{k}={v}" for k,v in mc.items()))
        elif mc: para(doc,f"다중 검정 보정: {mc}")
    if stt.get("software"):
        para(doc,f"분석 소프트웨어: {', '.join(stt['software'])}")
    if stt.get("extra"): para(doc,stt["extra"])
    # 표본 수
    para(doc,"표본 수 산출",bold=True)
    para(doc,c["sample_size"] or "[표본 수 미산출 — sample_size.py --from-prereg 결과를 narrative.sample_size에 기입]")

    # === 6. 연구대상자 안전보호 / 윤리 ===
    h1(doc,"6. 연구대상자 안전보호에 대한 대책 및 연구의 윤리성 확보를 위한 방안")
    render_ethics(doc,c)

    # === 7. 참고문헌 ===
    h1(doc,"7. 참고문헌")
    if c["references"]:
        for r in c["references"]:
            tag=[]
            if r.get("pmid"): tag.append(f"PMID:{r['pmid']}")
            if r.get("doi"): tag.append(f"doi:{r['doi']}")
            suffix=f"  ({'; '.join(tag)})" if tag else ""
            para(doc,f"{r.get('n','')}. {r.get('text','')}{suffix}")
    else:
        para(doc,"[참고문헌 없음 — search_log.json 또는 narrative.references 필요]")

    doc.save(outpath)

def render_ethics(doc,c):
    e=c.get("ethics",{}) or {}
    pi=c["pi"]
    defaults=[
        ("1. 연구 대상자 보호 원칙", e.get("protection") or
            "본 연구는 후향적/관찰적 자료 분석으로 대상자에 대한 추가적 침습·개입이 없다. "
            "연구 전 과정에서 대상자의 권리·안전·복지를 최우선으로 보호한다."),
        ("2. IRB 심의 분류 및 동의 면제", e.get("irb_class") or
            "「생명윤리 및 안전에 관한 법률」 및 기관 IRB 규정에 따라 심의 분류 및 동의 면제 여부를 IRB가 판단한다. "
            "후향적 의무기록 분석의 경우 동의 면제 요건(최소 위험, 면제 시 대상자 불이익 없음, 익명화) 충족 여부를 검토한다."),
        ("3. 데이터 익명화 및 개인정보 보호", e.get("privacy") or
            "분석 자료는 개인식별정보(PHI)를 제거한 익명화/가명화 자료를 사용하며, 접근은 사전 승인된 연구진에 한정한다."),
        ("4. 헬싱키 선언 준수", e.get("helsinki") or
            "본 연구는 헬싱키 선언(2024 개정)의 원칙에 따라 수행한다."),
        ("5. 자료 보관 및 폐기", e.get("retention") or
            "연구 종료 후 관련 법령에서 정한 기간(통상 3년) 동안 암호화 저장 후 안전하게 폐기한다."),
        ("6. AI 사용 공개 (ICMJE AI Disclosure)", e.get("ai_disclosure") or
            f"AI 기반 도구(Claude/Anthropic)를 가설 정제·분석 계획·문서 작성 보조에 사용하였다. "
            f"모든 학술적 판단과 책임은 저자({pi['name']})에게 귀속된다."),
        ("7. 이해상충 및 연구비", e.get("coi") or
            "모든 연구진은 연구 결과에 영향을 줄 수 있는 재정적·비재정적 이해관계가 없음을 확인한다."),
        ("8. 연구 결과 발표 계획", e.get("dissemination") or
            "SCI(E)급 국제 학술지 게재 및 학술대회 발표를 계획하며, 발표 시 개인식별정보를 완전히 제거한다."),
    ]
    for title,body in defaults:
        para(doc,title,bold=True); para(doc,body)

# ---------- IRB 메타데이터 ----------
def build_irb_metadata(prereg, content):
    return {
        "protocol_version": 1,
        "generated_at": datetime.datetime.now().astimezone().isoformat(),
        "linked_prereg_hash": prereg.get("hash"),
        "linked_prereg_version": prereg.get("version"),
        "study_type": content["study_type"],
        "irb_status": "pending_submission",
        "irb_number": None,
        "irb_approval_date": None,
        "irb_type": None,
        "irb_exempt_reason": None,
        "submission_log": [],
    }

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument("--prereg",required=True)
    ap.add_argument("--narrative")
    ap.add_argument("--profile")
    ap.add_argument("--refs")
    ap.add_argument("--outdir",default=".")
    ap.add_argument("--auto-sample-size",action="store_true",
                    help="narrative.sample_size 비었으면 sample_size.py로 자동 산출")
    a=ap.parse_args()
    prereg=load(a.prereg)
    if not prereg: sys.exit(f"prereg 로드 실패: {a.prereg}")
    narrative=load(a.narrative,{}) ; profile=load(a.profile,{}); refs=load(a.refs)
    content=resolve(prereg,narrative,profile,refs)

    # 표본 수 자동 산출
    if (a.auto_sample_size or not content["sample_size"]):
        try:
            sys.path.insert(0,os.path.dirname(os.path.abspath(__file__)))
            import sample_size as ss
            r=ss.from_prereg(a.prereg)
            if "error" not in r:
                content["sample_size"]=ss.korean_text(r)
            else:
                content.setdefault("_warnings",[]).append("sample_size: "+r["error"])
        except Exception as ex:
            content.setdefault("_warnings",[]).append(f"sample_size 자동산출 실패: {ex}")

    os.makedirs(a.outdir,exist_ok=True)
    docx_path=os.path.join(a.outdir,"research_protocol.docx")
    render(content,docx_path)
    with open(os.path.join(a.outdir,"protocol_content.resolved.json"),"w",encoding="utf-8") as f:
        json.dump(content,f,ensure_ascii=False,indent=2)
    with open(os.path.join(a.outdir,"irb_metadata.json"),"w",encoding="utf-8") as f:
        json.dump(build_irb_metadata(prereg,content),f,ensure_ascii=False,indent=2)
    print(json.dumps({"ok":True,"docx":docx_path,
                      "resolved":os.path.join(a.outdir,"protocol_content.resolved.json"),
                      "irb_metadata":os.path.join(a.outdir,"irb_metadata.json"),
                      "study_type":content["study_type"],
                      "warnings":content.get("_warnings",[])},ensure_ascii=False,indent=2))

if __name__=="__main__":
    main()
